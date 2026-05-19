import io
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Custom Numbered Canvas for PDF page counting ---
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            super().showPage()
        super().save()

    def draw_page_number(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#718096"))
        
        # Suppress footer on page 1 if desired, or show on all
        page_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(A4[0] - 54, 36, page_text)
        
        # Draw a thin footer separator line
        self.setStrokeColor(colors.HexColor("#E2E8F0"))
        self.setLineWidth(0.5)
        self.line(54, 48, A4[0] - 54, 48)
        
        # Add company tag in footer
        self.drawString(54, 36, "Livecode Technologies")
        self.restoreState()


def get_registration_details(registration, course=None):
    """Common utility to parse registration parameters."""
    first_name = registration.first_name
    last_name = registration.last_name
    org = registration.organization or "N/A"
    country = registration.country or "N/A"
    email = registration.email
    
    # Course Title
    course_title = course.title if course else (registration.course_title or "Master Business Analytics and Reporting with Microsoft Excel and Power BI Training Course")
    
    # Location
    location = registration.schedule_location or "N/A"
    if location == "N/A" and course and course.logistics:
        location = course.logistics.location or "Nairobi, Kenya"
    if not location or location == "N/A":
        location = "Nairobi, Kenya"

    # Date
    date_str = registration.schedule_date or "N/A"
    if date_str == "N/A" and course and course.logistics:
        start = course.logistics.start_date or ""
        end = course.logistics.end_date or ""
        if start and end:
            date_str = f"{start} - {end}"
        elif start:
            date_str = start
    if not date_str or date_str == "N/A":
        date_str = datetime.now().strftime("%B %d, %Y")

    # Duration
    duration = "10 days"
    if course and course.logistics and course.logistics.duration:
        duration = course.logistics.duration

    # Price
    price_usd = 2700.0
    if course and course.logistics and course.logistics.price_usd:
        price_usd = course.logistics.price_usd

    # Pax Count
    pax = 1
    if registration.registration_type == "group" and registration.group_size:
        try:
            pax = int(registration.group_size)
        except ValueError:
            pax = 1

    return {
        "first_name": first_name,
        "last_name": last_name,
        "org": org,
        "country": country,
        "email": email,
        "course_title": course_title,
        "location": location,
        "date_str": date_str,
        "duration": duration,
        "price_usd": price_usd,
        "pax": pax
    }


def generate_invoice_pdf(registration, course=None, group_members: list | None = None) -> io.BytesIO:
    """Dynamically generates a beautiful, 3-page professional invoice PDF using ReportLab."""
    details = get_registration_details(registration, course)
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=72
    )

    styles = getSampleStyleSheet()
    
    # Custom Brand Colors
    primary_color = colors.HexColor("#0F2942")  # Brand Navy
    secondary_color = colors.HexColor("#E28743")  # Brand Orange/Gold Accent
    text_dark = colors.HexColor("#2D3748")
    bg_light = colors.HexColor("#F7FAFC")
    border_color = colors.HexColor("#E2E8F0")

    # Custom Paragraph Styles
    title_style = ParagraphStyle(
        'InvoiceTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=primary_color,
        spaceAfter=15
    )
    
    h2_style = ParagraphStyle(
        'InvoiceHeading2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=primary_color,
        spaceBefore=10,
        spaceAfter=5
    )

    body_style = ParagraphStyle(
        'InvoiceBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=13,
        textColor=text_dark
    )

    body_bold = ParagraphStyle(
        'InvoiceBodyBold',
        parent=body_style,
        fontName='Helvetica-Bold'
    )

    header_cell_style = ParagraphStyle(
        'HeaderCell',
        parent=body_style,
        fontName='Helvetica-Bold',
        textColor=colors.white
    )

    story = []

    # ------------------ PAGE 1 ------------------
    # Top branding bar
    brand_table_data = [
        [
            Paragraph("<b>LIVECODE TECHNOLOGIES LTD</b><br/><font size=8 color='#718096'>Outreach & Corporate Training Division</font>", body_style),
            Paragraph("<b>TAX PIN:</b> 107267616<br/><b>REG NO:</b> 107267616", ParagraphStyle('TaxRight', parent=body_style, alignment=2))
        ]
    ]
    brand_table = Table(brand_table_data, colWidths=[300, 188])
    brand_table.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ('LINEBELOW', (0,0), (-1,-1), 1, border_color),
    ]))
    story.append(brand_table)
    story.append(Spacer(1, 20))

    # Title & Invoice details
    story.append(Paragraph("INVOICE", title_style))

    # Metadata Block (Attention to vs Invoice details)
    invoice_no = f"INV-{1000 + int(str(registration.id.int)[:6]) % 90000}"
    cust_id = f"CUST-{1000 + int(str(registration.id.int)[6:12]) % 9000}"
    date_today = datetime.now().strftime("%B %d, %Y")

    meta_left = [
        Paragraph(f"<b>Attention To:</b> {registration.title or 'Mr.'} {details['first_name']} {details['last_name']}", body_style),
        Paragraph(f"<b>Organization:</b> {details['org']}", body_style),
        Paragraph(f"<b>Country:</b> {details['country']}", body_style),
        Paragraph(f"<b>Email:</b> {details['email']}", body_style),
    ]
    
    meta_right = [
        Paragraph(f"<b>Invoice No:</b> {invoice_no}", body_style),
        Paragraph(f"<b>Date:</b> {date_today}", body_style),
        Paragraph(f"<b>Customer ID:</b> {cust_id}", body_style),
        Paragraph(f"<b>Payment Method:</b> Bank Transfer / Offline", body_style),
    ]

    meta_table_data = [
        [meta_left, meta_right]
    ]
    meta_table = Table(meta_table_data, colWidths=[244, 244])
    meta_table.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BACKGROUND', (0,0), (-1,-1), bg_light),
        ('BOX', (0,0), (-1,-1), 1, border_color),
        ('TOPPADDING', (0,0), (-1,-1), 12),
        ('BOTTOMPADDING', (0,0), (-1,-1), 12),
        ('LEFTPADDING', (0,0), (-1,-1), 12),
        ('RIGHTPADDING', (0,0), (-1,-1), 12),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 25))

    # Line Items Table
    price_val = details['price_usd']
    pax_val = details['pax']
    subtotal = price_val * pax_val
    tax = subtotal * 0.16
    total_amount = subtotal + tax

    pax_text = f"{pax_val} participant" if pax_val == 1 else f"{pax_val} participants"

    items_data = [
        [
            Paragraph("Project/Event Description", header_cell_style),
            Paragraph("No. of Pax", header_cell_style),
            Paragraph("Unit Price (USD)", header_cell_style),
            Paragraph("Total Price (USD)", header_cell_style)
        ],
        [
            Paragraph(f"<b>{details['course_title']}</b><br/><font size=8 color='#718096'>Location: {details['location']} | Dates: {details['date_str']}</font>", body_style),
            Paragraph(pax_text, body_style),
            Paragraph(f"{price_val:,.2f}", body_style),
            Paragraph(f"{subtotal:,.2f}", body_style)
        ],
        # Empty space or simple summaries
        ["", "", Paragraph("<b>SUBTOTAL:</b>", body_bold), Paragraph(f"USD {subtotal:,.2f}", body_bold)],
        ["", "", Paragraph("<b>TAX (16%):</b>", body_bold), Paragraph(f"USD {tax:,.2f}", body_bold)],
        ["", "", Paragraph("<b>TOTAL AMOUNT:</b>", ParagraphStyle('TotalBold', parent=body_bold, textColor=primary_color)), Paragraph(f"USD {total_amount:,.2f}", ParagraphStyle('TotalBold2', parent=body_bold, textColor=primary_color))]
    ]

    items_table = Table(items_data, colWidths=[240, 80, 84, 84])
    items_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), primary_color),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('ALIGN', (1,1), (-1,-1), 'RIGHT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('GRID', (0,0), (-1,1), 0.5, border_color),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ('LEFTPADDING', (0,0), (-1,-1), 10),
        ('RIGHTPADDING', (0,0), (-1,-1), 10),
        ('LINEABOVE', (2,2), (3,-1), 1, border_color),
    ]))
    story.append(items_table)
    story.append(Spacer(1, 20))

    # --- Participants List (group registrations) ---
    if group_members:
        participants_header = Paragraph("<b>GROUP PARTICIPANTS</b>", ParagraphStyle(
            'ParticipantsHead', parent=body_style, fontName='Helvetica-Bold',
            textColor=primary_color, fontSize=10
        ))
        story.append(participants_header)
        story.append(Spacer(1, 6))

        # Build lead row + member rows
        all_participants = [{
            "title": registration.title or "",
            "first_name": details["first_name"],
            "last_name": details["last_name"],
            "email": details["email"],
            "role": "Lead Registrant"
        }]
        for m in group_members:
            m_email = (m.get("email") or "").strip().lower()
            lead_email = (details["email"] or "").strip().lower()
            if m_email == lead_email:
                continue
            all_participants.append({
                "title": m.get("title", ""),
                "first_name": m.get("first_name", ""),
                "last_name": m.get("last_name", ""),
                "email": m.get("email", ""),
                "role": "Participant"
            })

        p_table_data = [
            [
                Paragraph("#", header_cell_style),
                Paragraph("Name", header_cell_style),
                Paragraph("Email", header_cell_style),
                Paragraph("Role", header_cell_style),
            ]
        ]
        for idx, p in enumerate(all_participants, start=1):
            full_name = f"{p['title']} {p['first_name']} {p['last_name']}".strip()
            p_table_data.append([
                Paragraph(str(idx), body_style),
                Paragraph(full_name, body_style),
                Paragraph(p["email"], body_style),
                Paragraph(p["role"], body_style),
            ])

        p_table = Table(p_table_data, colWidths=[24, 190, 190, 84])
        p_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), primary_color),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.4, border_color),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, bg_light]),
        ]))
        story.append(p_table)
        story.append(Spacer(1, 20))
    else:
        story.append(Spacer(1, 10))

    # Banking Details Card
    bank_data = [
        [Paragraph("<b>BANKING DETAILS FOR WIRE TRANSFER</b>", ParagraphStyle('BankHead', parent=body_style, fontName='Helvetica-Bold', textColor=primary_color)), ""],
        [Paragraph("<b>Name of Bank:</b>", body_style), Paragraph("Equity Bank Limited", body_style)],
        [Paragraph("<b>Branch Name:</b>", body_style), Paragraph("Harambee Avenue", body_style)],
        [Paragraph("<b>Account Name:</b>", body_style), Paragraph("Livecode Technologies Ltd", body_style)],
        [Paragraph("<b>Account Number:</b>", body_style), Paragraph("0240298633598", body_style)],
        [Paragraph("<b>Swift Code:</b>", body_style), Paragraph("EQBLKENA", body_style)],
        [Paragraph("<b>Bank Code:</b>", body_style), Paragraph("68", body_style)],
        [Paragraph("<b>Branch Code:</b>", body_style), Paragraph("024", body_style)]
    ]
    bank_table = Table(bank_data, colWidths=[150, 338])
    bank_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), bg_light),
        ('SPAN', (0,0), (1,0)),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('BOX', (0,0), (-1,-1), 1, border_color),
        ('LEFTPADDING', (0,0), (-1,-1), 15),
        ('RIGHTPADDING', (0,0), (-1,-1), 15),
        ('TOPPADDING', (0,0), (-1,0), 10),
        ('BOTTOMPADDING', (0,-1), (-1,-1), 10),
    ]))
    story.append(bank_table)
    story.append(Spacer(1, 20))
    story.append(Paragraph("<b>Note:</b> All checks should be payable to: <i>Livecode Technologies Ltd.</i><br/>Please send proof of payment to: <b>info@livecodetechnologies.com</b>", ParagraphStyle('NoteStyle', parent=body_style, fontSize=8.5)))

    story.append(PageBreak())

    # ------------------ PAGE 2 ------------------
    story.append(Paragraph("TERMS AND CONDITIONS", title_style))
    story.append(Spacer(1, 10))

    tc_data = [
        ("1. DEFINITIONS AND APPLICATION", "In these Terms and Conditions \"Livecode Technologies\" means Livecode Technologies Ltd. \"Client\" any person at whose request or on whose behalf Livecode Technologies undertakes any business or provides advice, information or services. If any legislation is compulsorily applicable to any business undertaken, these Terms and Conditions shall, as regards such business, be read as subject to such legislation and nothing in these Terms and Conditions shall be construed as a surrender by Livecode Technologies of any of its rights or immunities or as an increase of any of its responsibilities or liabilities under such legislation."),
        ("2. Validity of Offer", "Ninety days (90) from date of issue."),
        ("3. Pricing", "Training Fee shall be charged at individual or group rates which includes facilitation and cost of workshop materials. Prices are subject to change at any time prior to Livecode Technologies acceptance of Client’s order."),
        ("4. Order Acceptance", "Livecode Technologies shall accept the client order once the training fees have been paid in full."),
        ("5. Terms of Payment", "The client will make the full payment to the following account details: Livecode Technologies Ltd Account Number: 0240298633598 Equity Bank Limited, Harambee Avenue. Branch Code: 024 SWIFT Code: EQBLKENA. 16% VAT is payable until proof of exemption is provided."),
        ("6. Delivery Time", "Client-site workshop will have a minimum duration of 5 working days and is subject to availability of an instructor."),
        ("7. Quorum", "All scheduled courses will be conducted as long as there's a participant that has confirmed."),
        ("8. Force Majeure", "Neither party will be liable for performance delays, nor for non-performance due to causes beyond its reasonable control; however, will this provision not apply to Client's payment obligation.")
    ]

    for item_title, item_desc in tc_data:
        story.append(Paragraph(f"<b>{item_title}</b>", h2_style))
        story.append(Paragraph(item_desc, body_style))
        story.append(Spacer(1, 8))

    story.append(PageBreak())

    # ------------------ PAGE 3 ------------------
    story.append(Paragraph("TERMS AND CONDITIONS (CONTINUED)", title_style))
    story.append(Spacer(1, 10))

    tc_data_p3 = [
        ("9. Confidentiality", "Confidential information includes without limitation, pricing, software and hardware products, product plans, marketing and sales information, business plans, customer and supplier data, financial and technical information, \"know-how,\" trade secrets, and other information, whether such information is in written, oral, electronic, web-based, or other form. These may not be divulged to any third party without prior written consent from Livecode Technologies. The Client shall exercise the same degree of care as that used to protect their own confidential information but no less than reasonable care."),
        ("10. Jurisdiction and Law", "These Terms and Conditions and any act or contract to which they apply shall be governed by the Laws of Kenya and any dispute arising out of any act or contract to which these Terms and Conditions apply shall be subject to the exclusive jurisdiction of the Courts of Kenya."),
        ("11. Further Information", "Please contact us in case you need clarification or further information. Email: <b>info@livecodetechnologies.com</b>")
    ]

    for item_title, item_desc in tc_data_p3:
        story.append(Paragraph(f"<b>{item_title}</b>", h2_style))
        story.append(Paragraph(item_desc, body_style))
        story.append(Spacer(1, 8))

    story.append(Spacer(1, 30))
    story.append(Paragraph("<b>Acceptance Form</b>", ParagraphStyle('AccTitle', parent=styles['Heading2'], textColor=primary_color)))
    story.append(Spacer(1, 5))
    story.append(Paragraph("I hereby accept the terms and conditions of the above-mentioned invoice and place an order with Livecode Technologies Ltd for the services indicated on the invoice.", body_style))
    story.append(Spacer(1, 15))

    sig_data = [
        [Paragraph("<b>Company Name:</b> ___________________________", body_style), Paragraph("<b>Quote/Invoice No:</b> " + invoice_no, body_style)],
        [Paragraph("<b>Authorized Officer:</b> ________________________", body_style), Paragraph("<b>Date:</b> ___________________________", body_style)],
        [Paragraph("<b>Signature:</b> _________________________", body_style), Paragraph("<b>Official Stamp/Seal:</b> [ Seal Below ]", body_style)]
    ]
    sig_table = Table(sig_data, colWidths=[244, 244])
    sig_table.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('TOPPADDING', (0,0), (-1,-1), 10),
        ('BOTTOMPADDING', (0,0), (-1,-1), 10),
    ]))
    story.append(sig_table)

    # Build the document using the NumberedCanvas
    doc.build(story, canvasmaker=NumberedCanvas)
    buffer.seek(0)
    return buffer


def generate_invitation_letter_pdf(registration, course=None) -> io.BytesIO:
    """Generates a professional 6-page training invitation letter containing full curriculum blocks."""
    details = get_registration_details(registration, course)
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=72
    )

    styles = getSampleStyleSheet()
    primary_color = colors.HexColor("#0F2942")  # Brand Navy
    secondary_color = colors.HexColor("#E28743")
    text_dark = colors.HexColor("#2D3748")
    bg_light = colors.HexColor("#F7FAFC")
    border_color = colors.HexColor("#E2E8F0")

    title_style = ParagraphStyle(
        'LetterTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=primary_color,
        spaceAfter=15
    )

    h2_style = ParagraphStyle(
        'LetterH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=primary_color,
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'LetterBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=14,
        textColor=text_dark,
        spaceAfter=8
    )

    story = []

    # ------------------ PAGE 1: COVER/BRAND PAGE ------------------
    story.append(Spacer(1, 150))
    story.append(Paragraph("<font size=28 color='#0F2942'><b>OFFICIAL INVITATION LETTER</b></font>", ParagraphStyle('CoverTitle', parent=body_style, alignment=1)))
    story.append(Spacer(1, 15))
    story.append(Paragraph(f"<font size=14 color='#E28743'><b>{details['course_title']}</b></font>", ParagraphStyle('CoverSub', parent=body_style, alignment=1)))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"<font size=11 color='#718096'>Dates: {details['date_str']} | Location: {details['location']}</font>", ParagraphStyle('CoverMeta', parent=body_style, alignment=1)))
    story.append(Spacer(1, 200))
    
    org_info = """
    <b>LIVECODE TECHNOLOGIES LTD</b><br/>
    Town Office: Equity Plaza, Nairobi, Kenya<br/>
    Tel: +254 796 190 682 | Email: info@livecodetechnologies.com<br/>
    Web: www.livecodetechnologies.com
    """
    story.append(Paragraph(org_info, ParagraphStyle('CoverOrg', parent=body_style, alignment=1, fontSize=9, leading=14)))
    story.append(PageBreak())

    # ------------------ PAGE 2: FORMAL LETTER ------------------
    # Header block
    brand_hdr = [
        [
            Paragraph("<b>LIVECODE TECHNOLOGIES LTD</b>", body_style),
            Paragraph("<b>Date:</b> " + datetime.now().strftime("%d %B %Y"), ParagraphStyle('DateAl', parent=body_style, alignment=2))
        ]
    ]
    brand_tbl = Table(brand_hdr, colWidths=[244, 244])
    brand_tbl.setStyle(TableStyle([
        ('LINEBELOW', (0,0), (-1,-1), 0.5, border_color),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(brand_tbl)
    story.append(Spacer(1, 15))

    # Recipient info
    recipient_text = f"""
    To:<br/>
    <b>{details['org']}</b>,<br/>
    {details['country']}.<br/><br/>
    <b>Attn:</b> {registration.title or 'Mr.'} {details['first_name']} {details['last_name']}
    """
    story.append(Paragraph(recipient_text, body_style))
    story.append(Spacer(1, 10))

    # Subject line
    subj_text = f"<b>SUBJECT: INVITATION TO ATTEND A {details['duration'].upper()} TRAINING WORKSHOP ON {details['course_title'].upper()}</b>"
    story.append(Paragraph(subj_text, ParagraphStyle('Subj', parent=body_style, fontName='Helvetica-Bold', textColor=primary_color, fontSize=10, leading=14)))
    story.append(Spacer(1, 12))

    # Body paragraphs
    intro = f"""We are delighted to invite you to participate in a {details['duration']} training workshop <b>"{details['course_title']}"</b> to be held from <b>{details['date_str']}</b> in <b>{details['location']}</b>."""
    story.append(Paragraph(intro, body_style))

    overview = """
    <b>Course Overview</b><br/>
    Take your analysis and reporting skills to an advanced professional level. This course combines standard analytical best practices with the power of modern business intelligence tools. Learn pivots, imports, data formatting, Dax expressions, dashboards, and automated custom reporting.
    """
    story.append(Paragraph(overview, body_style))

    audience = """
    <b>Target Audience</b><br/>
    This program is designed for business analysts, operations supervisors, IT coordinators, operations assistants, finance assistants, marketing professionals, small business owners, and administrators who want to master dynamic data reports.
    """
    story.append(Paragraph(audience, body_style))

    story.append(Spacer(1, 20))
    story.append(Paragraph("Yours Sincerely,", body_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("<b>Managing Director</b><br/>Livecode Technologies Ltd", ParagraphStyle('CEOStyle', parent=body_style, leading=13)))

    story.append(PageBreak())

    # ------------------ PAGES 3 - 5: COURSE OUTLINE / SYLLABUS ------------------
    story.append(Paragraph("COURSE OUTLINE & SYLLABUS", title_style))
    story.append(Spacer(1, 10))

    outline_data = [
        ("Course Objectives", [
            "Gain enhanced spreadsheets data skills: creating templates, charts, graphics, and formulas for business and operational work.",
            "How to calculate with advanced functions & formulas, create interactive charts using Excel.",
            "How to analyze data using Pivot Charts, insert graphic objects and quickly transform business data into informative reports.",
            "How to use business dashboards to compare multiple elements in one screen.",
            "How to generate management and weekly reports quickly and accurately."
        ]),
        ("Module 1: Introduction to Business Data Analysis with Excel", [
            "Functions and Formulae (Entering and Copying Formulae)",
            "Creating Formulas using Names (Managing Names using Name Manager)",
            "Naming Ranges (Creating Named Ranges, Inserting Names in Formulae)",
            "Formatting (Number Formatting, Cell Styles, Clearing Formatting)",
            "Case Study/Practical Component: Analyzing a real-world business dataset."
        ]),
        ("Module 2: Calculating with Formulas and Functions", [
            "Use Date and Time Functions",
            "Use Statistical Functions (Database Functions)",
            "Use Lookup Functions (Logical Functions)",
            "Case Study/Practical Component: Calculating key business metrics."
        ]),
        ("Module 3: Visualizing Data with Charts", [
            "Use Recommended Charts Tool (Create Chart with Quick Analysis)",
            "Align Chart Axis Labels (Display Value Axis in Millions)",
            "Create Combo Chart (Create Chart Template for Sharing)",
            "Build Dynamic Chart with Table",
            "Case Study/Practical Component: Creating a comprehensive data visualization report."
        ]),
        ("Module 4: Analyzing Data with Tables", [
            "Create Table with Quick Analysis (Remove Duplicate Records)",
            "Add Total Row to Table (Insert Calculated Columns)",
            "Filter Table Data with Custom Filter (Use Slicer to Filter Data)",
            "Apply Conditional Formatting (Sort and Filter Data by Color)",
            "Case Study/Practical Component: Managing and analyzing a large dataset."
        ]),
        ("Module 5: Analyzing Data with PivotTables", [
            "Overview of PivotTable & PivotChart",
            "Create PivotTable from Table (Change Report Layouts)",
            "Show and Hide Report Totals (Create Custom Calculations)",
            "Filter Pivot Data with Slicer and Timeline",
            "Case Study/Practical Component: Developing a sales report using PivotTables."
        ]),
        ("Module 6: Importing External Data Sources", [
            "Import Data from Text Files",
            "Create PivotTable & PivotChart from Microsoft Access",
            "Manage Data Connections",
            "Case Study/Practical Component: Integrating and analyzing data from multiple external sources."
        ]),
        ("Module 7: Data Manipulation in Excel", [
            "How Excel Handles Different Data Types",
            "Data Consistency (Building Datasheets)",
            "Sorting and Filtering",
            "Data Cleaning and Formatting",
            "Case Study/Practical Component: Cleaning a messy dataset for business analysis."
        ]),
        ("Module 8: Creating Dashboard Reports", [
            "Overview of Dashboard Architecture",
            "Create PivotTable and PivotChart",
            "Hide Field Buttons from PivotChart",
            "Use Slicer to Connect PivotTables (Build Slicer Dashboard Report)",
            "Case Study/Practical Component: Building an interactive KPI Excel dashboard."
        ]),
        ("Module 9: Power BI Essentials", [
            "Power BI Overview & Setting Up",
            "Installing Power BI Desktop",
            "Connecting to Multiple Data Sources (Data Modelling)",
            "Introduction to DAX (Crossjoin, CALCULATE, VALUES, CALENDAR)",
            "Case Study/Practical Component: Creating a data model in Power BI."
        ]),
        ("Module 10: Business Data Analytics with Power BI", [
            "Connecting to and Cleaning Data",
            "Creating Calculated Fields & Measures",
            "Building Charts (Line, Bar, Geo Maps)",
            "Creating a Dashboard & Slicers",
            "Case Study/Practical Component: Developing a full interactive Power BI dashboard."
        ])
    ]

    # Render syllabus pages
    for idx, (section_title, bullet_items) in enumerate(outline_data):
        # We can put page breaks between modules if it gets too crowded
        if idx in [4, 8]:
            story.append(PageBreak())
            story.append(Paragraph("COURSE OUTLINE (CONTINUED)", title_style))
            story.append(Spacer(1, 10))

        story.append(Paragraph(section_title, h2_style))
        for bullet in bullet_items:
            story.append(Paragraph(f"• {bullet}", ParagraphStyle('BulletStyle', parent=body_style, leftIndent=12, firstLineIndent=-8)))
        story.append(Spacer(1, 6))

    story.append(PageBreak())

    # ------------------ PAGE 6: METHODOLOGY & LOGISTICS ------------------
    story.append(Paragraph("METHODOLOGY & ACCREDITATION", title_style))
    story.append(Spacer(1, 10))

    methodology = """
    <b>Training Methodology</b><br/>
    The instructor-led training is delivered using a blended learning approach and comprises of presentations, guided sessions of practical exercises, web-based templates and intensive group case studies. Facilitators are industry experts with years of practical experience. All facilitation and materials are offered in English.
    """
    story.append(Paragraph(methodology, body_style))
    story.append(Spacer(1, 10))

    accreditation = """
    <b>Accreditation & Certification</b><br/>
    Upon successful completion of this training, participants will be issued with a Livecode Technologies official training certificate.
    """
    story.append(Paragraph(accreditation, body_style))
    story.append(Spacer(1, 10))

    venue = f"""
    <b>Training Venue</b><br/>
    The training will be conducted in <b>{details['location']}</b>. The course registration fee covers the course tuition, extensive training materials, certificates, break refreshments, and daily lunches. Travel, insurance, hotel accommodation and personal expenses are separately catered for by the participant.
    """
    story.append(Paragraph(venue, body_style))
    story.append(Spacer(1, 10))

    inquiries = """
    <b>Inquiries & Custom Options</b><br/>
    This course can also be fully customized/tailored for your organization. For custom deliveries or inquiries, contact us on Tel: <b>+254 796 190 682</b> or email <b>info@livecodetechnologies.com</b>.
    """
    story.append(Paragraph(inquiries, body_style))

    doc.build(story, canvasmaker=NumberedCanvas)
    buffer.seek(0)
    return buffer


def generate_pre_training_form_docx(registration, course=None) -> io.BytesIO:
    """Generates a professional table-based Microsoft Word (.docx) document for pre-training evaluation."""
    details = get_registration_details(registration, course)
    
    doc = docx.Document()
    
    # Custom XML helpers to add cell background styling for table headers
    def set_cell_background(cell, hex_color):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Font Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Header
    title = doc.add_paragraph()
    title_run = title.add_run("PRE-TRAINING ASSESSMENT & EVALUATION FORM")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = docx.shared.RGBColor(15, 41, 66)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Livecode Technologies Ltd")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)

    # Instruction Note
    note = doc.add_paragraph()
    note_run = note.add_run("Instruction: To serve you best and tailor the training course outline, please take time to complete this questionnaire prior to the start of the session. Return this completed document to: info@livecodetechnologies.com.")
    note_run.italic = True
    note_run.font.size = Pt(9.5)
    note.paragraph_format.space_after = Pt(15)

    # Participant Details Table
    doc.add_heading("1. Participant & Contact Details", level=2)
    p_table = doc.add_table(rows=7, cols=2)
    p_table.style = 'Table Grid'
    
    p_data = [
        ("Course Title", details['course_title']),
        ("Course Dates", details['date_str']),
        ("Venue / Location", details['location']),
        ("Participant Name", f"{registration.title or 'Mr.'} {details['first_name']} {details['last_name']}"),
        ("Organisation & Dept", f"{details['org']} | {registration.department or 'N/A'}"),
        ("Email Address", details['email']),
        ("WhatsApp/Tel Number", registration.phone or "N/A")
    ]
    
    for idx, (label, val) in enumerate(p_data):
        row = p_table.rows[idx]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(val)
        
        # Style cell widths
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.8)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # Section 1 Questionnaire
    doc.add_heading("2. Background and Motivation", level=2)
    
    q_table = doc.add_table(rows=4, cols=2)
    q_table.style = 'Table Grid'
    
    # Header row
    hdr_row = q_table.rows[0]
    hdr_row.cells[0].paragraphs[0].add_run("Evaluation Aspect").bold = True
    hdr_row.cells[1].paragraphs[0].add_run("Participant's Response / Remarks").bold = True
    for cell in hdr_row.cells:
        set_cell_background(cell, "0F2942")
        cell.paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(255, 255, 255)
    
    q_items_s1 = [
        ("Why do you want to attend this course?", "[ ] Improve skills at current job\n[ ] Working on a specific upcoming project\n[ ] Career transition / promotion\n[ ] Other: ___________________________"),
        ("What is your current Educational background & Years of Experience?", "Degree/Qualifications:\n\nYears of Experience:"),
        ("What specific type of work do you currently carry out?", "\n\n")
    ]
    
    for idx, (label, default_val) in enumerate(q_items_s1):
        row = q_table.rows[idx+1]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(default_val)
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(4.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # Section 2 Questionnaire
    doc.add_heading("3. Data Operations & Expectations", level=2)
    
    q_table_s2 = doc.add_table(rows=5, cols=2)
    q_table_s2.style = 'Table Grid'
    
    # Header row
    hdr_row2 = q_table_s2.rows[0]
    hdr_row2.cells[0].paragraphs[0].add_run("Evaluation Aspect").bold = True
    hdr_row2.cells[1].paragraphs[0].add_run("Participant's Response / Remarks").bold = True
    for cell in hdr_row2.cells:
        set_cell_background(cell, "0F2942")
        cell.paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(255, 255, 255)
        
    q_items_s2 = [
        ("What type of data or reports do you regularly work with?", "\n\n"),
        ("What specific data challenges do you face in your job?", "\n\n"),
        ("What applications or tools (Excel, Power BI, SQL, etc.) do you want to master by the end?", "\n\n"),
        ("Which specific modules in the outline do you think will be most valuable?", "\n\n")
    ]
    
    for idx, (label, default_val) in enumerate(q_items_s2):
        row = q_table_s2.rows[idx+1]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(default_val)
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(4.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # Section 3 Questionnaire
    doc.add_heading("4. Level Rating & Comments", level=2)
    
    q_table_s3 = doc.add_table(rows=4, cols=2)
    q_table_s3.style = 'Table Grid'
    
    # Header row
    hdr_row3 = q_table_s3.rows[0]
    hdr_row3.cells[0].paragraphs[0].add_run("Evaluation Aspect").bold = True
    hdr_row3.cells[1].paragraphs[0].add_run("Participant's Response / Remarks").bold = True
    for cell in hdr_row3.cells:
        set_cell_background(cell, "0F2942")
        cell.paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(255, 255, 255)
        
    q_items_s3 = [
        ("Apart from the syllabus, list any custom areas you would like included:", "\n\n"),
        ("Rate your current knowledge level in this training field:", "[ ] Beginner (Level 1)\n[ ] Intermediate (Level 2)\n[ ] Advanced (Level 3)"),
        ("Any additional remarks, comments, or dietary preferences:", "\n\n")
    ]
    
    for idx, (label, default_val) in enumerate(q_items_s3):
        row = q_table_s3.rows[idx+1]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(default_val)
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(4.0)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
