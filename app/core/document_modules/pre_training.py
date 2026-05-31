import io

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.core.document_modules.common import get_registration_details


def generate_pre_training_form_docx(registration, course=None) -> io.BytesIO:
    """Generates a professional table-based Microsoft Word (.docx) document for pre-training evaluation."""
    details = get_registration_details(registration, course)
    
    doc = docx.Document()
    
    # Custom XML helpers to add cell background styling for table headers
    def set_cell_background(cell, hex_color):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Font Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Header
    title = doc.add_paragraph()
    title_run = title.add_run("PRE-TRAINING ASSESSMENT & EVALUATION FORM")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = docx.shared.RGBColor(15, 41, 66)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Livecode Technologies Ltd")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)

    # Instruction Note
    note = doc.add_paragraph()
    note_run = note.add_run("Instruction: To serve you best and tailor the training course outline, please take time to complete this questionnaire prior to the start of the session. Return this completed document to: info@livecodetechnologies.com.")
    note_run.italic = True
    note_run.font.size = Pt(9.5)
    note.paragraph_format.space_after = Pt(15)

    # Participant Details Table
    doc.add_heading("1. Participant & Contact Details", level=2)
    p_table = doc.add_table(rows=7, cols=2)
    p_table.style = 'Table Grid'
    
    p_data = [
        ("Course Title", details['course_title']),
        ("Course Dates", details['date_str']),
        ("Venue / Location", details['location']),
        ("Participant Name", f"{registration.title or 'Mr.'} {details['first_name']} {details['last_name']}"),
        ("Organisation & Dept", f"{details['org']} | {registration.department or 'N/A'}"),
        ("Email Address", details['email']),
        ("WhatsApp/Tel Number", registration.phone or "N/A")
    ]
    
    for idx, (label, val) in enumerate(p_data):
        row = p_table.rows[idx]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(val)
        
        # Style cell widths
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.8)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # Section 1 Questionnaire
    doc.add_heading("2. Background and Motivation", level=2)
    
    q_table = doc.add_table(rows=4, cols=2)
    q_table.style = 'Table Grid'
    
    # Header row
    hdr_row = q_table.rows[0]
    hdr_row.cells[0].paragraphs[0].add_run("Evaluation Aspect").bold = True
    hdr_row.cells[1].paragraphs[0].add_run("Participant's Response / Remarks").bold = True
    for cell in hdr_row.cells:
        set_cell_background(cell, "0F2942")
        cell.paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(255, 255, 255)
    
    q_items_s1 = [
        ("Why do you want to attend this course?", "[ ] Improve skills at current job\n[ ] Working on a specific upcoming project\n[ ] Career transition / promotion\n[ ] Other: ___________________________"),
        ("What is your current Educational background & Years of Experience?", "Degree/Qualifications:\n\nYears of Experience:"),
        ("What specific type of work do you currently carry out?", "\n\n")
    ]
    
    for idx, (label, default_val) in enumerate(q_items_s1):
        row = q_table.rows[idx+1]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(default_val)
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(4.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # Section 2 Questionnaire
    doc.add_heading("3. Data Operations & Expectations", level=2)
    
    q_table_s2 = doc.add_table(rows=5, cols=2)
    q_table_s2.style = 'Table Grid'
    
    # Header row
    hdr_row2 = q_table_s2.rows[0]
    hdr_row2.cells[0].paragraphs[0].add_run("Evaluation Aspect").bold = True
    hdr_row2.cells[1].paragraphs[0].add_run("Participant's Response / Remarks").bold = True
    for cell in hdr_row2.cells:
        set_cell_background(cell, "0F2942")
        cell.paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(255, 255, 255)
        
    q_items_s2 = [
        ("What type of data or reports do you regularly work with?", "\n\n"),
        ("What specific data challenges do you face in your job?", "\n\n"),
        ("What applications or tools (Excel, Power BI, SQL, etc.) do you want to master by the end?", "\n\n"),
        ("Which specific modules in the outline do you think will be most valuable?", "\n\n")
    ]
    
    for idx, (label, default_val) in enumerate(q_items_s2):
        row = q_table_s2.rows[idx+1]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(default_val)
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(4.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # Section 3 Questionnaire
    doc.add_heading("4. Level Rating & Comments", level=2)
    
    q_table_s3 = doc.add_table(rows=4, cols=2)
    q_table_s3.style = 'Table Grid'
    
    # Header row
    hdr_row3 = q_table_s3.rows[0]
    hdr_row3.cells[0].paragraphs[0].add_run("Evaluation Aspect").bold = True
    hdr_row3.cells[1].paragraphs[0].add_run("Participant's Response / Remarks").bold = True
    for cell in hdr_row3.cells:
        set_cell_background(cell, "0F2942")
        cell.paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(255, 255, 255)
        
    q_items_s3 = [
        ("Apart from the syllabus, list any custom areas you would like included:", "\n\n"),
        ("Rate your current knowledge level in this training field:", "[ ] Beginner (Level 1)\n[ ] Intermediate (Level 2)\n[ ] Advanced (Level 3)"),
        ("Any additional remarks, comments, or dietary preferences:", "\n\n")
    ]
    
    for idx, (label, default_val) in enumerate(q_items_s3):
        row = q_table_s3.rows[idx+1]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(default_val)
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(4.0)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
